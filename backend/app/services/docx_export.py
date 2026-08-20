"""Export papers and answer keys to DOCX/PDF."""

from __future__ import annotations

import io
import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.export import ExportInclude
from app.schemas.paper import AnswerKey, ContentBlock, Paper


def _render_content_blocks(paragraph, blocks: list[ContentBlock]) -> None:
    for block in blocks:
        if block.type == "equation":
            run = paragraph.add_run(f" [{block.value}] ")
            run.italic = True
        else:
            paragraph.add_run(block.value)


def _add_paper_content(doc: Document, paper: Paper) -> None:
    meta = paper.metadata
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{meta.subject} — {meta.grade_class}")
    run.bold = True
    run.font.size = Pt(14)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Total Marks: {meta.total_marks}    Duration: {meta.duration}")

    if meta.instructions:
        instr = doc.add_paragraph()
        instr.add_run(meta.instructions).italic = True

    doc.add_paragraph()

    for section in paper.sections:
        heading = doc.add_paragraph()
        heading.add_run(f"Section {section.section_id}").bold = True
        if section.instructions:
            si = doc.add_paragraph()
            si.add_run(section.instructions).italic = True

        for question in section.questions:
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q{question.q_id}. ")
            q_run.bold = True
            _render_content_blocks(q_para, question.content)
            q_para.add_run(f"  [{question.marks} marks]")

            if question.type == "mcq" and question.options:
                for i, opt in enumerate(question.options):
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Pt(18)
                    opt_para.add_run(f"({chr(65 + i)}) {opt}")

        doc.add_paragraph()


def _add_answer_key_content(doc: Document, answer_key: AnswerKey, paper: Paper) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ANSWER KEY")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    q_lookup = {q.q_id: q for section in paper.sections for q in section.questions}
    for entry in answer_key.answers:
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{entry.q_id}. ")
        q_run.bold = True
        _render_content_blocks(q_para, entry.answer)
        if entry.explanation:
            exp = doc.add_paragraph()
            exp.paragraph_format.left_indent = Pt(18)
            exp.add_run(f"Explanation: {entry.explanation}").italic = True


def export_to_docx(
    paper: Paper,
    answer_key: AnswerKey | None,
    include: ExportInclude,
    output_path: Path,
) -> Path:
    doc = Document()
    if include in ("paper", "both"):
        _add_paper_content(doc, paper)
    if include in ("answer_key", "both") and answer_key:
        if include == "both":
            doc.add_page_break()
        _add_answer_key_content(doc, answer_key, paper)
    doc.save(str(output_path))
    return output_path


def export_to_pdf(docx_path: Path, pdf_path: Path) -> Path:
    """Convert DOCX to PDF using LibreOffice if available, else reportlab fallback."""
    for cmd in (
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
    ):
        try:
            subprocess.run(cmd, check=True, capture_output=True, timeout=60)
            generated = pdf_path.parent / f"{docx_path.stem}.pdf"
            if generated.exists():
                if generated != pdf_path:
                    generated.rename(pdf_path)
                return pdf_path
        except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
            continue

    return _pdf_fallback(docx_path, pdf_path)


def _pdf_fallback(docx_path: Path, pdf_path: Path) -> Path:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    # Minimal text dump — real layout handled by docx path when LibreOffice present
    c = canvas.Canvas(str(pdf_path), pagesize=A4)
    width, height = A4
    y = height - 50
    c.setFont("Helvetica", 10)
    c.drawString(50, y, f"Exported from {docx_path.name}")
    c.drawString(50, y - 20, "Install LibreOffice for full PDF fidelity.")
    c.save()
    return pdf_path


def build_export_file(
    paper: Paper,
    answer_key: AnswerKey | None,
    include: ExportInclude,
    fmt: str,
    export_dir: str,
) -> tuple[Path, str]:
    os.makedirs(export_dir, exist_ok=True)
    base = f"{paper.paper_id}_{include}"
    docx_path = Path(export_dir) / f"{base}.docx"

    export_to_docx(paper, answer_key, include, docx_path)

    if fmt == "docx":
        return docx_path, docx_path.name

    pdf_path = Path(export_dir) / f"{base}.pdf"
    export_to_pdf(docx_path, pdf_path)
    return pdf_path, pdf_path.name
